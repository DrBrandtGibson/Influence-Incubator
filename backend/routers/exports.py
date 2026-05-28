"""PDF + DOCX export endpoints for plans."""
from io import BytesIO
from typing import Dict, List, Any, Optional
import json

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, Color
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, KeepTogether
)
from reportlab.pdfgen import canvas as pdfcanvas

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth_supabase import require_user, anon_client_with_token, CurrentUser
from access import has_pro_access

router = APIRouter(prefix="/plans", tags=["exports"])

# ----------------------- Brand palette -----------------------
BRAND_CHARCOAL = HexColor("#14100C")
BRAND_GOLD = HexColor("#D4AF37")
BRAND_BRONZE = HexColor("#8B6F2A")
BRAND_CREAM = HexColor("#F8F2E5")
BRAND_MUTED = HexColor("#6B6359")

STEP_TITLES = {
    1: ("DEFINE", "Your Purpose"),
    2: ("EXTRACT", "Your Audience"),
    3: ("FRAME", "Your Story"),
    4: ("IGNITE", "Your Brand"),
    5: ("NURTURE", "The Transformation"),
    6: ("EXPAND", "Your Influence"),
    7: ("DELIVER", "Exceptional Service"),
}


def _safe_json(raw):
    if not raw:
        return None
    try:
        return json.loads(raw)
    except Exception:
        return None


def _by_key(inputs: List[dict]) -> Dict[int, Dict[str, str]]:
    """Group plan_inputs into {step_num: {field_key: value}}."""
    out: Dict[int, Dict[str, str]] = {}
    for row in inputs:
        sn = row.get("step_num")
        if not sn:
            continue
        out.setdefault(sn, {})[row.get("field_key")] = row.get("value") or ""
    return out


# Map of (step_num, field_key) -> human label, for primary fields we want to surface.
LABELS: Dict[int, List[tuple]] = {
    1: [
        ("business_name", "Business Name"),
        ("mtp_statement", "Massive Transformative Purpose"),
        ("fp_q5", "Purpose Statement"),
        ("why_level_7", "Deep WHY"),
        ("chief_q3_what", "3-Month Goal (WHAT)"),
        ("chief_y1_what", "1-Year Goal (WHAT)"),
        ("chief_y3_what", "3-Year Goal (WHAT)"),
        ("chief_y5_what", "5-Year Goal (WHAT)"),
        ("structure_chosen", "Business Structure"),
    ],
    2: [
        ("dc_name", "Dream Customer"),
        ("niche_type", "Niche Type"),
        ("micro_niche_statement", "Micro-Niche Statement"),
    ],
    3: [
        ("brand_voice_statement", "Brand Voice"),
        ("hero_journey_founder_narration", "Founder's Journey · AI Narration"),
        ("hero_journey_customer_narration", "Customer's Journey · AI Narration"),
    ],
    5: [
        ("cp_name", "Continuity Program"),
        ("cp_price", "Continuity Price"),
        ("cp_what_monthly", "Continuity — Monthly Value"),
    ],
}

# ----------------------- Framework constants (mirror frontend/lib/framework.js) -----------------------
# Kept in sync with framework.js to render full Step 2/3/4 details in exports.
_MASLOW_LABELS = {
    "physiological": "Physiological",
    "safety": "Safety",
    "love_belonging": "Love & Belonging",
    "esteem": "Esteem",
    "self_actualization": "Self-Actualization",
    "transcendence": "Transcendence",
}
_NEED_LABELS = {
    "certainty": "Certainty",
    "variety": "Variety",
    "significance": "Significance",
    "love_connection": "Love / Connection",
    "growth": "Growth",
    "contribution": "Contribution",
}
_DEMO_Q = [
    ("age", "Age range"),
    ("gender", "Gender / identity"),
    ("location", "Location"),
    ("income", "Income / wealth"),
    ("education", "Education"),
    ("occupation", "Occupation"),
    ("family", "Family situation"),
    ("lifestyle", "Lifestyle markers"),
]
_PSYCHO_Q = [
    ("values", "Core values"),
    ("beliefs", "Beliefs"),
    ("aspirations", "Aspirations"),
    ("fears", "Fears"),
    ("daily_habits", "Daily habits"),
    ("influences", "Influences"),
    ("frustrations", "Frustrations"),
    ("vocabulary", "Vocabulary / how they talk"),
]
_STORY_BANK_CATEGORIES = [
    ("origin", "Origin"),
    ("rupture", "Rupture / Wake-up"),
    ("mentor", "Mentor"),
    ("transformation", "Transformation"),
    ("failure", "Failure"),
    ("resurrection", "Resurrection"),
    ("calling", "Calling"),
    ("teacher", "Teaching Moment"),
    ("future", "Future Vision"),
]
_POCKET_MEDIA_CHANNELS = [
    ("newsletter", "Newsletter"),
    ("blog", "Blog"),
    ("podcast", "Podcast"),
    ("video", "Video"),
    ("events", "Events"),
]
_POCKET_MEDIA_FIELDS = [
    ("name", "Working Name"),
    ("url", "URL"),
    ("format", "Format"),
    ("cadence", "Cadence"),
    ("audience_pull", "Audience Pull"),
    ("kpi", "KPI"),
    ("first_5_ideas", "First 5 Ideas"),
]
_WEBSITE_PAGES = {
    "influencer": [("home", "Home"), ("about", "About / Story"), ("work_with", "Work With Me"), ("podcast_or_blog", "Podcast / Blog"), ("free_resource", "Free Resource"), ("speak", "Speaking"), ("contact", "Contact")],
    "medical":    [("home", "Home"), ("about", "About Provider"), ("services", "Services"), ("patient_resources", "Patient Resources"), ("team", "Team"), ("locations", "Locations & Hours"), ("book_appt", "Book Appointment"), ("contact", "Contact")],
}
_MARKETING_TRACKS = [("diy", "DIY Track"), ("ai10x", "10X-with-AI Track")]
_MARKETING_FIELDS = [
    ("weekly_schedule", "Weekly Schedule"),
    ("tools", "Tools / Stack"),
    ("time_investment", "Time Investment"),
    ("expected_outcome", "Expected Outcome"),
]
_CAL_PHASES = [("d30", "Days 0–30"), ("d60", "Days 31–60"), ("d90", "Days 61–90"), ("beyond", "Beyond 90")]
_CAL_PILLARS = [("content", "Content"), ("engagement", "Engagement"), ("growth", "Growth"), ("offer", "Offer")]


def _watermark(canvas, doc):
    """Diagonal watermark drawn on every page (free tier only)."""
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 70)
    canvas.setFillColor(Color(0.85, 0.85, 0.85, alpha=0.30))
    canvas.translate(LETTER[0] / 2, LETTER[1] / 2)
    canvas.rotate(45)
    canvas.drawCentredString(0, 0, "FREE PREVIEW")
    canvas.setFont("Helvetica", 14)
    canvas.drawCentredString(0, -50, "Upgrade to Pro for clean export")
    canvas.restoreState()


def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(BRAND_MUTED)
    canvas.drawString(0.6 * inch, 0.4 * inch, "The Influence Incubator Formula · A 7-Step Plan")
    canvas.drawRightString(LETTER[0] - 0.6 * inch, 0.4 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _on_page_free(canvas, doc):
    _watermark(canvas, doc)
    _footer(canvas, doc)


def _on_page_pro(canvas, doc):
    _footer(canvas, doc)


def _build_styles():
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=28, textColor=BRAND_CHARCOAL, spaceAfter=8, leading=32),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"], fontName="Helvetica-Oblique", fontSize=13, textColor=BRAND_BRONZE, spaceAfter=18),
        "eyebrow": ParagraphStyle("eyebrow", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8, textColor=BRAND_BRONZE, spaceAfter=4, leading=10),
        "stepHeading": ParagraphStyle("stepHeading", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=22, textColor=BRAND_CHARCOAL, spaceBefore=4, spaceAfter=10, leading=26),
        "sectionHeading": ParagraphStyle("sectionHeading", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=BRAND_CHARCOAL, spaceBefore=10, spaceAfter=4, leading=16),
        "label": ParagraphStyle("label", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=9, textColor=BRAND_BRONZE, spaceAfter=2, leading=11),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica", fontSize=10.5, textColor=BRAND_CHARCOAL, leading=15, spaceAfter=6),
        "italic": ParagraphStyle("italic", parent=base["Normal"], fontName="Helvetica-Oblique", fontSize=10.5, textColor=BRAND_MUTED, leading=15, spaceAfter=6),
        "muted": ParagraphStyle("muted", parent=base["Normal"], fontName="Helvetica", fontSize=9.5, textColor=BRAND_MUTED, leading=13, spaceAfter=4),
    }
    return styles


def _para(text, style):
    """HTML-escape and create a Paragraph. Returns Spacer for empty text."""
    if not text:
        return Spacer(1, 0)
    safe = (str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>"))
    return Paragraph(safe, style)


def _pdf_cover(plan: dict, is_pro: bool, styles, story: list) -> None:
    story.append(_para("THE INFLUENCE INCUBATOR FORMULA", styles["eyebrow"]))
    story.append(_para(plan.get("title") or "Untitled Plan", styles["title"]))
    story.append(_para("A 7-Step Business Plan", styles["subtitle"]))
    if not is_pro:
        story.append(_para("Free preview — upgrade to Pro for the full export without watermark and to enable Word format.", styles["italic"]))
    story.append(Spacer(1, 14))


def _pdf_step_header(step_num: int, styles, story: list) -> None:
    verb, tail = STEP_TITLES[step_num]
    story.append(PageBreak())
    story.append(_para(f"STEP 0{step_num}", styles["eyebrow"]))
    story.append(_para(f"{verb} <font color='#8B6F2A'>{tail}</font>", styles["stepHeading"]))


def _pdf_labeled_fields(step_num: int, step_data: dict, styles, story: list) -> None:
    for fk, label in LABELS.get(step_num, []):
        v = step_data.get(fk)
        if v:
            story.append(_para(label.upper(), styles["label"]))
            story.append(_para(v, styles["body"]))


# Dispatch table for step-specific structured renderings.
_PDF_STEP_RENDERERS = {
    2: lambda d, s, st: _render_step2(d, s, st),
    3: lambda d, s, st: _render_step3(d, s, st),
    4: lambda d, s, st: _render_step4(d, s, st),
    5: lambda d, s, st: _render_step5(d, s, st),
    6: lambda d, s, st: _render_step6(d, s, st),
    7: lambda d, s, st: _render_step7(d, s, st),
}


def _build_pdf(plan: dict, inputs: List[dict], is_pro: bool) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.8 * inch,
        title=plan.get("title") or "Plan",
        author="Influence Incubator Formula",
    )
    styles = _build_styles()
    data = _by_key(inputs)
    story: list = []

    _pdf_cover(plan, is_pro, styles, story)

    # Free users only get Steps 1 & 2 in the export, regardless of what's filled in.
    step_range = range(1, 8) if is_pro else range(1, 3)
    for step_num in step_range:
        _pdf_step_header(step_num, styles, story)
        step_data = data.get(step_num, {})
        if not step_data:
            story.append(_para("No entries yet for this step.", styles["italic"]))
            continue
        _pdf_labeled_fields(step_num, step_data, styles, story)
        renderer = _PDF_STEP_RENDERERS.get(step_num)
        if renderer:
            renderer(step_data, story, styles)

    # Footer for free users: list locked steps and upgrade CTA
    if not is_pro:
        story.append(PageBreak())
        story.append(_para("STEPS 3 – 7 LOCKED", styles["eyebrow"]))
        story.append(_para("Upgrade to Pro to export the full 7-step plan", styles["stepHeading"]))
        story.append(_para(
            "Your free export includes Steps 1 (DEFINE) and 2 (EXTRACT). The remaining five steps "
            "— FRAME, IGNITE, NURTURE, EXPAND, and DELIVER — are part of Pro. Visit your dashboard "
            "and choose Lifetime ($97 one-time) or Monthly ($19/mo) to unlock everything, including "
            "the editable Word export.",
            styles["body"]))

    on_page = _on_page_free if not is_pro else _on_page_pro
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return buf.getvalue()


def _render_step2(d: Dict[str, str], story, styles):
    """Dream Customer details: demographics + psychographics + Maslow + 6 Needs."""
    # Demographics
    demo_rows = [(label, d.get(f"demo_{k}")) for k, label in _DEMO_Q if d.get(f"demo_{k}")]
    if demo_rows:
        story.append(Spacer(1, 8))
        story.append(_para("DEMOGRAPHICS", styles["label"]))
        for label, v in demo_rows:
            story.append(_para(f"<b>{label}:</b> {v}", styles["body"]))
    # Psychographics
    psycho_rows = [(label, d.get(f"psycho_{k}")) for k, label in _PSYCHO_Q if d.get(f"psycho_{k}")]
    if psycho_rows:
        story.append(Spacer(1, 6))
        story.append(_para("PSYCHOGRAPHICS", styles["label"]))
        for label, v in psycho_rows:
            story.append(_para(f"<b>{label}:</b> {v}", styles["body"]))
    # Maslow's
    try:
        ml = json.loads(d.get("maslow_levels") or "[]")
    except Exception:
        ml = []
    ml_labels = [_MASLOW_LABELS.get(k, k) for k in ml if k]
    if ml_labels:
        story.append(Spacer(1, 6))
        story.append(_para("MASLOW'S HIERARCHY (selected)", styles["label"]))
        story.append(_para(" · ".join(ml_labels), styles["body"]))
    # 6 Core Needs
    try:
        nd = json.loads(d.get("robbins_needs") or "[]")
    except Exception:
        nd = []
    nd_labels = [_NEED_LABELS.get(k, k) for k in nd if k]
    if nd_labels:
        story.append(Spacer(1, 6))
        story.append(_para("6 CORE HUMAN NEEDS (selected)", styles["label"]))
        story.append(_para(" · ".join(nd_labels), styles["body"]))


def _render_step3(d: Dict[str, str], story, styles):
    """Frame Your Story: journey narrations + offers (HSO + Important Stories) + Story Bank."""
    # Founder & Customer Journey narrations
    founder_nar = d.get("hero_journey_founder_narration")
    customer_nar = d.get("hero_journey_customer_narration")
    if founder_nar:
        story.append(Spacer(1, 8))
        story.append(_para("FOUNDER'S JOURNEY · AI SYNTHESIS", styles["label"]))
        story.append(_para(founder_nar, styles["body"]))
    if customer_nar:
        story.append(Spacer(1, 4))
        story.append(_para("CUSTOMER'S JOURNEY · AI SYNTHESIS", styles["label"]))
        story.append(_para(customer_nar, styles["body"]))

    # Offers
    try:
        offer_ids = json.loads(d.get("offer_ids") or "[]")
    except Exception:
        offer_ids = []
    if offer_ids:
        story.append(Spacer(1, 8))
        story.append(_para("OFFERS", styles["label"]))
        for i, oid in enumerate(offer_ids):
            name = d.get(f"{oid}_name") or f"Offer {i + 1}"
            price = d.get(f"{oid}_price") or ""
            rationale = d.get(f"{oid}_price_rationale")
            hook = d.get(f"{oid}_hook")
            stry = d.get(f"{oid}_story")
            promise = d.get(f"{oid}_promise")
            elevator = d.get(f"{oid}_elevator")
            story.append(_para(f"<b>{name}</b>" + (f" — <font color='#8B6F2A'>{price}</font>" if price else ""), styles["sectionHeading"]))
            if rationale:
                story.append(_para(f"<i>Pricing rationale:</i> {rationale}", styles["body"]))
            # Stack
            try:
                stack = json.loads(d.get(f"{oid}_stack") or "[]")
            except Exception:
                stack = []
            stack = [r for r in stack if isinstance(r, dict) and (r.get("item") or r.get("benefit") or r.get("value"))]
            if stack:
                story.append(_para("<i>Stack:</i>", styles["body"]))
                for r in stack:
                    parts = [f"<b>{r.get('item', '')}</b>"]
                    if r.get("benefit"):
                        parts.append(f"— {r['benefit']}")
                    if r.get("value"):
                        parts.append(f"· <font color='#8B6F2A'>{r['value']}</font>")
                    story.append(_para("&nbsp;&nbsp;• " + " ".join(parts), styles["body"]))
            if hook:
                story.append(_para(f"<i>Hook:</i> {hook}", styles["body"]))
            if stry:
                story.append(_para(f"<i>Story:</i> {stry}", styles["body"]))
            if promise:
                story.append(_para(f"<i>Transformation Promise:</i> {promise}", styles["body"]))
            if elevator:
                story.append(_para(f"<i>Important Story — Elevator Pitch:</i> {elevator}", styles["body"]))

    # Story Bank — 9 sections
    sb_entries = []
    for cat_key, cat_label in _STORY_BANK_CATEGORIES:
        # Collect all keys starting with cat_key_ from d
        cat_items = [(fk, v) for fk, v in d.items() if fk.startswith(f"{cat_key}_") and v]
        if cat_items:
            sb_entries.append((cat_label, cat_items))
    if sb_entries:
        story.append(Spacer(1, 10))
        story.append(_para("STORY BANK · FULL ARCHIVE", styles["label"]))
        for cat_label, items in sb_entries:
            story.append(_para(f"<b>{cat_label}</b>", styles["sectionHeading"]))
            for _, v in items:
                story.append(_para(f"• {v}", styles["body"]))


def _render_step4(d: Dict[str, str], story, styles):
    """IGNITE Your Brand Card: archetype + palette w/ HEX + 3-level typography +
    full channel cards + drafted pages + marketing plan + 30/60/90 calendar."""
    primary = d.get("archetype_primary")
    secondary = d.get("archetype_secondary")
    if primary or secondary:
        story.append(_para("ARCHETYPE", styles["label"]))
        s = primary or "—"
        if secondary:
            s += f" · with {secondary}"
        story.append(_para(s, styles["body"]))

    # Palette + HEX codes
    palette_chosen = d.get("palette_chosen")
    palette = None
    try:
        palettes = (json.loads(d.get("palettes_json") or "{}") or {}).get("palettes") or []
        palette = next((p for p in palettes if p.get("name") == palette_chosen), None)
    except Exception:
        pass
    if palette_chosen:
        story.append(Spacer(1, 4))
        story.append(_para("PALETTE", styles["label"]))
        story.append(_para(f"<b>{palette_chosen}</b>" + (f" — <i>{palette.get('mood', '')}</i>" if palette and palette.get("mood") else ""), styles["body"]))
        if palette and palette.get("colors"):
            hex_str = "  ·  ".join([f"<font color='{c}'>■</font> <font face='Courier'>{c}</font>" for c in palette["colors"]])
            story.append(_para(hex_str, styles["body"]))

    # Typography (3 levels)
    typo_chosen = d.get("typography_chosen")
    typo = None
    try:
        pairings = (json.loads(d.get("typography_json") or "{}") or {}).get("pairings") or []
        typo = next((p for p in pairings if p.get("name") == typo_chosen), None)
    except Exception:
        pass
    if typo_chosen:
        story.append(Spacer(1, 4))
        story.append(_para("TYPOGRAPHY", styles["label"]))
        story.append(_para(f"<b>{typo_chosen}</b>", styles["body"]))
        if typo:
            headline = typo.get("headline") or typo.get("heading") or "—"
            subhead = typo.get("subheadline") or typo.get("heading") or headline
            body = typo.get("body") or "—"
            story.append(_para(f"Headline: <b>{headline}</b> · Subheadline: <b>{subhead}</b> · Body: <b>{body}</b>", styles["body"]))

    # Pocket Media — full channel cards
    enabled_channels = [(k, label) for k, label in _POCKET_MEDIA_CHANNELS if d.get(f"pm_{k}_enabled") == "yes"]
    if enabled_channels:
        story.append(Spacer(1, 8))
        story.append(_para("POCKET MEDIA EMPIRE · CHANNELS IN THE MIX", styles["label"]))
        for k, label in enabled_channels:
            story.append(_para(f"<b>{label}</b>", styles["sectionHeading"]))
            for fk, flabel in _POCKET_MEDIA_FIELDS:
                v = d.get(f"pm_{k}_{fk}")
                if v:
                    story.append(_para(f"<i>{flabel}:</i> {v}", styles["body"]))

    # Website — drafted pages
    site_tpl = d.get("site_template")
    if site_tpl:
        pages = _WEBSITE_PAGES.get(site_tpl, [])
        drafted = [(pk, plabel, d.get(f"site_{site_tpl}_{pk}")) for pk, plabel in pages if d.get(f"site_{site_tpl}_{pk}")]
        story.append(Spacer(1, 6))
        story.append(_para(f"WEBSITE — {site_tpl.upper()} HUB", styles["label"]))
        if drafted:
            for _, plabel, v in drafted:
                story.append(_para(f"<b>{plabel}</b>", styles["sectionHeading"]))
                story.append(_para(v, styles["body"]))
        else:
            story.append(_para("(No pages drafted yet.)", styles["italic"]))

    # Marketing Plan — both tracks
    has_any_mt = any(d.get(f"mt_{tk}_{fk}") for tk, _ in _MARKETING_TRACKS for fk, _ in _MARKETING_FIELDS)
    if has_any_mt:
        story.append(Spacer(1, 6))
        story.append(_para("MARKETING PLAN · BOTH TRACKS", styles["label"]))
        for tk, tlabel in _MARKETING_TRACKS:
            rows = [(flabel, d.get(f"mt_{tk}_{fk}")) for fk, flabel in _MARKETING_FIELDS if d.get(f"mt_{tk}_{fk}")]
            if not rows:
                continue
            story.append(_para(f"<b>{tlabel}</b>", styles["sectionHeading"]))
            for flabel, v in rows:
                story.append(_para(f"<i>{flabel}:</i> {v}", styles["body"]))

    # 30/60/90 + Beyond Calendar
    has_any_cal = any(d.get(f"cal_{pk}_{phk}") for pk, _ in _CAL_PILLARS for phk, _ in _CAL_PHASES)
    if has_any_cal:
        story.append(Spacer(1, 6))
        story.append(_para("30/60/90 + BEYOND CONTENT CALENDAR", styles["label"]))
        for pk, plabel in _CAL_PILLARS:
            cells = [(phlabel, d.get(f"cal_{pk}_{phk}")) for phk, phlabel in _CAL_PHASES if d.get(f"cal_{pk}_{phk}")]
            if not cells:
                continue
            story.append(_para(f"<b>{plabel}</b>", styles["sectionHeading"]))
            for phlabel, v in cells:
                story.append(_para(f"<i>{phlabel}:</i> {v}", styles["body"]))


def _render_step5(d: Dict[str, str], story, styles):
    """Framework, SaaS, Community."""
    fw = _safe_json(d.get("framework_json"))
    if fw:
        story.append(_para("TRANSFORMATIVE FRAMEWORK", styles["label"]))
        story.append(_para(f"<b>{fw.get('name', '—')}</b> — <i>{fw.get('tagline', '')}</i>", styles["body"]))
        for i, ph in enumerate(fw.get("phases") or []):
            line = f"{i + 1}. <b>{ph.get('verb', '')} {ph.get('name', '')}</b>"
            if ph.get("transformation"):
                line += f" — <i>{ph['transformation']}</i>"
            if ph.get("description"):
                line += f"<br/>&nbsp;&nbsp;&nbsp;{ph['description']}"
            story.append(_para(line, styles["body"]))
    saas = _safe_json(d.get("saas_options_json"))
    if saas and saas.get("options"):
        story.append(_para("SAAS OPPORTUNITIES", styles["label"]))
        for o in saas["options"]:
            story.append(_para(f"<b>{o.get('name', 'Option')}</b> — {o.get('problem', '')}", styles["body"]))
    com = _safe_json(d.get("community_json"))
    if com:
        story.append(_para("COMMUNITY", styles["label"]))
        story.append(_para(f"<b>{com.get('name', '—')}</b> — {com.get('member_archetype', '')}", styles["body"]))


def _render_step6(d: Dict[str, str], story, styles):
    """Dream 100 + Live Events list + Book outline."""
    try:
        d100 = json.loads(d.get("dream100_list") or "[]")
    except Exception:
        d100 = []
    if d100:
        story.append(_para(f"DREAM 100 — {len(d100)} entries", styles["label"]))

    # Events
    try:
        event_ids = json.loads(d.get("events_ids") or "[]")
    except Exception:
        event_ids = []
    if event_ids:
        story.append(Spacer(1, 6))
        story.append(_para(f"LIVE EVENTS & CHALLENGES — {len(event_ids)}", styles["label"]))
        for i, eid in enumerate(event_ids):
            tk = d.get(f"{eid}_type")
            name = d.get(f"{eid}_name") or f"Event {i + 1}"
            promise = d.get(f"{eid}_promise")
            outcome = d.get(f"{eid}_outcome")
            conv = d.get(f"{eid}_conversion")
            head = name + (f" — <i>{tk}</i>" if tk else "")
            story.append(_para(f"<b>{head}</b>", styles["sectionHeading"]))
            if promise:
                story.append(_para(f"<i>Promise:</i> {promise}", styles["body"]))
            if outcome:
                story.append(_para(f"<i>Outcome:</i> {outcome}", styles["body"]))
            if conv:
                story.append(_para(f"<i>Conversion path:</i> {conv}", styles["body"]))

    outline = _safe_json(d.get("book_outline_json"))
    if outline:
        story.append(Spacer(1, 6))
        story.append(_para("BOOK", styles["label"]))
        story.append(_para(f"<b>{outline.get('title', '—')}</b> — <i>{outline.get('subtitle', '')}</i>", styles["body"]))
        chapters = outline.get("chapters") or []
        if chapters:
            story.append(_para(f"{len(chapters)} chapters outlined", styles["muted"]))
            for ch in chapters[:30]:
                story.append(_para(f"&nbsp;&nbsp;{ch.get('n', '')}. {ch.get('title', '')}", styles["body"]))


def _render_step7(d: Dict[str, str], story, styles):
    """DELIVER Exceptional Service Card — journey map + onboarding + S&D + quality + feedback."""
    journey = _safe_json(d.get("journey_json"))
    stages = (journey or {}).get("stages") or []
    if stages:
        story.append(Spacer(1, 6))
        story.append(_para("CUSTOMER JOURNEY MAP", styles["label"]))
        for i, s in enumerate(stages):
            story.append(_para(f"<b>{i + 1}. {s.get('name', 'Stage')}</b>", styles["sectionHeading"]))
            if s.get("customer_does"):
                story.append(_para(f"<i>Customer does:</i> {s['customer_does']}", styles["body"]))
            if s.get("customer_feels"):
                story.append(_para(f"<i>Customer feels:</i> {s['customer_feels']}", styles["body"]))
            if s.get("we_do"):
                story.append(_para(f"<i>We do:</i> {s['we_do']}", styles["body"]))
            if s.get("risk"):
                story.append(_para(f"<i>Risk:</i> {s['risk']}", styles["body"]))

    onb = _safe_json(d.get("onboarding_json"))
    seq = (onb or {}).get("sequence") or []
    if seq:
        story.append(Spacer(1, 6))
        story.append(_para(f"ONBOARDING SEQUENCE — {len(seq)} TOUCHPOINTS", styles["label"]))
        for i, s in enumerate(seq):
            head = f"D{s.get('day')}" if s.get("day") else f"#{i + 1}"
            title = s.get("title") or s.get("action") or ""
            story.append(_para(f"<b>{head} · {title}</b>" + (f"  ({s['channel']})" if s.get("channel") else ""), styles["sectionHeading"]))
            if s.get("detail"):
                story.append(_para(s["detail"], styles["body"]))

    ret = _safe_json(d.get("retention_json"))
    plays = (ret or {}).get("plays") or []
    if plays:
        story.append(Spacer(1, 6))
        story.append(_para(f"SURPRISE & DELIGHT PLAYBOOK — {len(plays)} PLAYS", styles["label"]))
        for p in plays:
            head = p.get("name") or p.get("trigger") or "Play"
            story.append(_para(f"<b>{head}</b>", styles["sectionHeading"]))
            if p.get("trigger"):
                story.append(_para(f"<i>Trigger:</i> {p['trigger']}", styles["body"]))
            if p.get("gesture"):
                story.append(_para(f"<i>Gesture:</i> {p['gesture']}", styles["body"]))
            if p.get("cost"):
                story.append(_para(f"<i>Cost:</i> {p['cost']}", styles["body"]))

    quality_keys = [("dq_response_sla", "Response SLA"), ("dq_quality_bar", "Quality bar"), ("dq_audit", "Audit/QA cadence"), ("dq_recovery", "Recovery protocol")]
    quality_rows = [(label, d.get(k)) for k, label in quality_keys if d.get(k)]
    if quality_rows:
        story.append(Spacer(1, 6))
        story.append(_para("QUALITY STANDARDS", styles["label"]))
        for label, v in quality_rows:
            story.append(_para(f"<b>{label}:</b> {v}", styles["body"]))

    feedback_keys = [("df_nps_cadence", "NPS / Pulse cadence"), ("df_listening", "Listening surface"), ("df_close", "Close-the-loop"), ("df_signal", "Strongest signal to act on")]
    feedback_rows = [(label, d.get(k)) for k, label in feedback_keys if d.get(k)]
    if feedback_rows:
        story.append(Spacer(1, 6))
        story.append(_para("FEEDBACK LOOP", styles["label"]))
        for label, v in feedback_rows:
            story.append(_para(f"<b>{label}:</b> {v}", styles["body"]))

    # AI-generated 10-step Do-This-Next list (if generated on Business Plan page)
    todos_json = _safe_json(d.get("business_plan_todos_json"))
    todos = (todos_json or {}).get("todos") or []
    if todos:
        story.append(Spacer(1, 8))
        story.append(_para("DO THIS NEXT — AI-PRIORITIZED 10-STEP LIST", styles["label"]))
        for i, t in enumerate(todos[:10]):
            story.append(_para(f"<b>{i + 1}. {t.get('title', '—')}</b>", styles["sectionHeading"]))
            if t.get("rationale"):
                story.append(_para(t["rationale"], styles["body"]))


# =============================== DOCX ===============================

_BRONZE = RGBColor(0x8B, 0x6F, 0x2A)


def _docx_styled_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size_pt: Optional[int] = None, color=None):
    """Append a styled run to a docx paragraph and return it."""
    r = paragraph.add_run(text)
    r.bold = bool(bold)
    r.italic = bool(italic)
    if size_pt:
        r.font.size = Pt(size_pt)
    if color is not None:
        r.font.color.rgb = color
    return r


def _docx_set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.8)


def _docx_cover(doc: Document, plan: dict) -> None:
    _docx_styled_run(doc.add_paragraph(), "THE INFLUENCE INCUBATOR FORMULA", bold=True, size_pt=9, color=_BRONZE)
    _docx_styled_run(doc.add_paragraph(), plan.get("title") or "Untitled Plan", bold=True, size_pt=28)
    _docx_styled_run(doc.add_paragraph(), "A 7-Step Business Plan", italic=True, size_pt=13, color=_BRONZE)


def _docx_step_header(doc: Document, step_num: int) -> None:
    verb, tail = STEP_TITLES[step_num]
    _docx_styled_run(doc.add_paragraph(), f"STEP 0{step_num}", bold=True, size_pt=8, color=_BRONZE)
    heading = doc.add_paragraph()
    _docx_styled_run(heading, f"{verb} ", bold=True, size_pt=22)
    _docx_styled_run(heading, tail, bold=True, size_pt=22, color=_BRONZE)


def _docx_labeled_fields(doc: Document, step_num: int, step_data: dict) -> None:
    for fk, label in LABELS.get(step_num, []):
        v = step_data.get(fk)
        if not v:
            continue
        _docx_styled_run(doc.add_paragraph(), label.upper(), bold=True, size_pt=9, color=_BRONZE)
        doc.add_paragraph(v)


def _build_docx(plan: dict, inputs: List[dict]) -> bytes:
    doc = Document()
    _docx_set_margins(doc)
    data = _by_key(inputs)

    _docx_cover(doc, plan)
    for step_num in range(1, 8):
        doc.add_page_break()
        _docx_step_header(doc, step_num)
        step_data = data.get(step_num, {})
        if not step_data:
            doc.add_paragraph("No entries yet for this step.").italic = True
            continue
        _docx_labeled_fields(doc, step_num, step_data)
        _docx_step_extras(doc, step_num, step_data)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_render_step2(doc: Document, d: Dict[str, str]) -> None:
    """Dream Customer details for Word export."""
    demo_rows = [(label, d.get(f"demo_{k}")) for k, label in _DEMO_Q if d.get(f"demo_{k}")]
    if demo_rows:
        _docx_styled_run(doc.add_paragraph(), "Demographics", bold=True, size_pt=11, color=_BRONZE)
        for label, v in demo_rows:
            p = doc.add_paragraph()
            _docx_styled_run(p, f"{label}: ", bold=True)
            p.add_run(v)
    psycho_rows = [(label, d.get(f"psycho_{k}")) for k, label in _PSYCHO_Q if d.get(f"psycho_{k}")]
    if psycho_rows:
        _docx_styled_run(doc.add_paragraph(), "Psychographics", bold=True, size_pt=11, color=_BRONZE)
        for label, v in psycho_rows:
            p = doc.add_paragraph()
            _docx_styled_run(p, f"{label}: ", bold=True)
            p.add_run(v)
    try:
        ml = json.loads(d.get("maslow_levels") or "[]")
    except Exception:
        ml = []
    ml_labels = [_MASLOW_LABELS.get(k, k) for k in ml if k]
    if ml_labels:
        _docx_styled_run(doc.add_paragraph(), "Maslow's Hierarchy (selected)", bold=True, size_pt=11, color=_BRONZE)
        doc.add_paragraph(" · ".join(ml_labels))
    try:
        nd = json.loads(d.get("robbins_needs") or "[]")
    except Exception:
        nd = []
    nd_labels = [_NEED_LABELS.get(k, k) for k in nd if k]
    if nd_labels:
        _docx_styled_run(doc.add_paragraph(), "6 Core Human Needs (selected)", bold=True, size_pt=11, color=_BRONZE)
        doc.add_paragraph(" · ".join(nd_labels))


def _docx_render_step3(doc: Document, d: Dict[str, str]) -> None:
    """Step 3 — Journey narrations + offers HSO bundle + Story Bank."""
    # Narrations are rendered via _docx_labeled_fields already (we added to LABELS).
    try:
        offer_ids = json.loads(d.get("offer_ids") or "[]")
    except Exception:
        offer_ids = []
    for i, oid in enumerate(offer_ids):
        name = d.get(f"{oid}_name") or f"Offer {i + 1}"
        price = d.get(f"{oid}_price") or ""
        rationale = d.get(f"{oid}_price_rationale")
        p = doc.add_paragraph()
        _docx_styled_run(p, name, bold=True, size_pt=13)
        if price:
            _docx_styled_run(p, f"  —  {price}", color=_BRONZE)
        if rationale:
            pr = doc.add_paragraph()
            _docx_styled_run(pr, "Pricing rationale: ", bold=True)
            pr.add_run(rationale)
        # Stack
        try:
            stack = json.loads(d.get(f"{oid}_stack") or "[]")
        except Exception:
            stack = []
        stack = [r for r in stack if isinstance(r, dict) and (r.get("item") or r.get("benefit") or r.get("value"))]
        if stack:
            _docx_styled_run(doc.add_paragraph(), "Stack:", bold=True)
            for r in stack:
                line = f"• {r.get('item', '')}"
                if r.get("benefit"):
                    line += f" — {r['benefit']}"
                if r.get("value"):
                    line += f" · {r['value']}"
                doc.add_paragraph(line)
        for key, label in (("hook", "Hook"), ("story", "Story"), ("promise", "Transformation Promise"), ("elevator", "Important Story — Elevator Pitch")):
            v = d.get(f"{oid}_{key}")
            if not v:
                continue
            pp = doc.add_paragraph()
            _docx_styled_run(pp, f"{label}: ", bold=True)
            pp.add_run(v)

    # Story Bank — 9 sections
    for cat_key, cat_label in _STORY_BANK_CATEGORIES:
        cat_items = [v for fk, v in d.items() if fk.startswith(f"{cat_key}_") and v]
        if not cat_items:
            continue
        _docx_styled_run(doc.add_paragraph(), f"Story Bank · {cat_label}", bold=True, size_pt=11, color=_BRONZE)
        for v in cat_items:
            doc.add_paragraph(f"• {v}")


def _docx_render_step4(doc: Document, d: Dict[str, str]) -> None:
    """Full IGNITE Your Brand Card for Word export."""
    primary = d.get("archetype_primary")
    secondary = d.get("archetype_secondary")
    if primary:
        p = doc.add_paragraph()
        _docx_styled_run(p, "Archetype: ", bold=True)
        p.add_run(primary + (f" · with {secondary}" if secondary else ""))

    # Palette
    palette_chosen = d.get("palette_chosen")
    if palette_chosen:
        try:
            palettes = (json.loads(d.get("palettes_json") or "{}") or {}).get("palettes") or []
            palette = next((p for p in palettes if p.get("name") == palette_chosen), None)
        except Exception:
            palette = None
        p = doc.add_paragraph()
        _docx_styled_run(p, "Palette: ", bold=True)
        p.add_run(palette_chosen)
        if palette and palette.get("colors"):
            doc.add_paragraph(" · ".join(palette["colors"]))

    # Typography
    typo_chosen = d.get("typography_chosen")
    if typo_chosen:
        try:
            pairings = (json.loads(d.get("typography_json") or "{}") or {}).get("pairings") or []
            typo = next((p for p in pairings if p.get("name") == typo_chosen), None)
        except Exception:
            typo = None
        p = doc.add_paragraph()
        _docx_styled_run(p, "Typography: ", bold=True)
        p.add_run(typo_chosen)
        if typo:
            headline = typo.get("headline") or typo.get("heading") or "—"
            subhead = typo.get("subheadline") or typo.get("heading") or headline
            body = typo.get("body") or "—"
            doc.add_paragraph(f"Headline: {headline} · Subheadline: {subhead} · Body: {body}")

    # Pocket Media — channels
    enabled = [(k, label) for k, label in _POCKET_MEDIA_CHANNELS if d.get(f"pm_{k}_enabled") == "yes"]
    if enabled:
        _docx_styled_run(doc.add_paragraph(), "Pocket Media Empire — Channels in the Mix", bold=True, size_pt=12, color=_BRONZE)
        for k, label in enabled:
            _docx_styled_run(doc.add_paragraph(), label, bold=True, size_pt=11)
            for fk, flabel in _POCKET_MEDIA_FIELDS:
                v = d.get(f"pm_{k}_{fk}")
                if v:
                    pp = doc.add_paragraph()
                    _docx_styled_run(pp, f"{flabel}: ", bold=True)
                    pp.add_run(v)

    # Website pages
    site_tpl = d.get("site_template")
    if site_tpl:
        pages = _WEBSITE_PAGES.get(site_tpl, [])
        drafted = [(plabel, d.get(f"site_{site_tpl}_{pk}")) for pk, plabel in pages if d.get(f"site_{site_tpl}_{pk}")]
        if drafted:
            _docx_styled_run(doc.add_paragraph(), f"Website — {site_tpl} Hub", bold=True, size_pt=12, color=_BRONZE)
            for plabel, v in drafted:
                _docx_styled_run(doc.add_paragraph(), plabel, bold=True, size_pt=11)
                doc.add_paragraph(v)

    # Marketing Plan
    for tk, tlabel in _MARKETING_TRACKS:
        rows = [(flabel, d.get(f"mt_{tk}_{fk}")) for fk, flabel in _MARKETING_FIELDS if d.get(f"mt_{tk}_{fk}")]
        if not rows:
            continue
        _docx_styled_run(doc.add_paragraph(), f"Marketing Plan — {tlabel}", bold=True, size_pt=12, color=_BRONZE)
        for flabel, v in rows:
            pp = doc.add_paragraph()
            _docx_styled_run(pp, f"{flabel}: ", bold=True)
            pp.add_run(v)

    # 30/60/90 Calendar
    any_cal = any(d.get(f"cal_{pk}_{phk}") for pk, _ in _CAL_PILLARS for phk, _ in _CAL_PHASES)
    if any_cal:
        _docx_styled_run(doc.add_paragraph(), "30/60/90 + Beyond Calendar", bold=True, size_pt=12, color=_BRONZE)
        for pk, plabel in _CAL_PILLARS:
            cells = [(phlabel, d.get(f"cal_{pk}_{phk}")) for phk, phlabel in _CAL_PHASES if d.get(f"cal_{pk}_{phk}")]
            if not cells:
                continue
            _docx_styled_run(doc.add_paragraph(), plabel, bold=True, size_pt=11)
            for phlabel, v in cells:
                pp = doc.add_paragraph()
                _docx_styled_run(pp, f"{phlabel}: ", bold=True)
                pp.add_run(v)


def _docx_render_step5(doc: Document, d: Dict[str, str]) -> None:
    fw = _safe_json(d.get("framework_json"))
    if not fw or not fw.get("phases"):
        return
    for i, ph in enumerate(fw["phases"]):
        doc.add_paragraph(f"{i + 1}. {ph.get('verb', '')} {ph.get('name', '')} — {ph.get('transformation', '')}")


def _docx_render_step6(doc: Document, d: Dict[str, str]) -> None:
    """Step 6 — Live Events + Book outline."""
    try:
        event_ids = json.loads(d.get("events_ids") or "[]")
    except Exception:
        event_ids = []
    if event_ids:
        _docx_styled_run(doc.add_paragraph(), f"Live Events & Challenges ({len(event_ids)})", bold=True, size_pt=12, color=_BRONZE)
        for i, eid in enumerate(event_ids):
            tk = d.get(f"{eid}_type")
            name = d.get(f"{eid}_name") or f"Event {i + 1}"
            promise = d.get(f"{eid}_promise")
            outcome = d.get(f"{eid}_outcome")
            conv = d.get(f"{eid}_conversion")
            p = doc.add_paragraph()
            _docx_styled_run(p, name, bold=True, size_pt=11)
            if tk:
                _docx_styled_run(p, f"  ({tk})", italic=True)
            for label, v in (("Promise", promise), ("Outcome", outcome), ("Conversion path", conv)):
                if not v:
                    continue
                pp = doc.add_paragraph()
                _docx_styled_run(pp, f"{label}: ", bold=True)
                pp.add_run(v)
    outline = _safe_json(d.get("book_outline_json"))
    if outline and outline.get("chapters"):
        _docx_styled_run(doc.add_paragraph(), f"Book — {outline.get('title') or '—'}", bold=True, size_pt=12, color=_BRONZE)
        for ch in outline["chapters"]:
            doc.add_paragraph(f"{ch.get('n', '')}. {ch.get('title', '')}")


def _docx_render_step7(doc: Document, d: Dict[str, str]) -> None:
    """DELIVER Exceptional Service Card."""
    journey = _safe_json(d.get("journey_json"))
    stages = (journey or {}).get("stages") or []
    if stages:
        _docx_styled_run(doc.add_paragraph(), "Customer Journey Map", bold=True, size_pt=12, color=_BRONZE)
        for i, s in enumerate(stages):
            _docx_styled_run(doc.add_paragraph(), f"{i + 1}. {s.get('name', 'Stage')}", bold=True, size_pt=11)
            for label, key in (("Customer does", "customer_does"), ("Customer feels", "customer_feels"), ("We do", "we_do"), ("Risk", "risk")):
                v = s.get(key)
                if v:
                    pp = doc.add_paragraph()
                    _docx_styled_run(pp, f"{label}: ", bold=True)
                    pp.add_run(v)
    onb = _safe_json(d.get("onboarding_json"))
    seq = (onb or {}).get("sequence") or []
    if seq:
        _docx_styled_run(doc.add_paragraph(), f"Onboarding Sequence — {len(seq)} touchpoints", bold=True, size_pt=12, color=_BRONZE)
        for i, s in enumerate(seq):
            head = f"D{s.get('day')}" if s.get("day") else f"#{i + 1}"
            _docx_styled_run(doc.add_paragraph(), f"{head}  ·  {s.get('title') or s.get('action') or ''}", bold=True, size_pt=11)
            if s.get("detail"):
                doc.add_paragraph(s["detail"])
    ret = _safe_json(d.get("retention_json"))
    plays = (ret or {}).get("plays") or []
    if plays:
        _docx_styled_run(doc.add_paragraph(), f"Surprise & Delight Playbook — {len(plays)} plays", bold=True, size_pt=12, color=_BRONZE)
        for p in plays:
            _docx_styled_run(doc.add_paragraph(), p.get("name") or p.get("trigger") or "Play", bold=True, size_pt=11)
            for label, key in (("Trigger", "trigger"), ("Gesture", "gesture"), ("Cost", "cost")):
                v = p.get(key)
                if v:
                    pp = doc.add_paragraph()
                    _docx_styled_run(pp, f"{label}: ", bold=True)
                    pp.add_run(v)
    # Quality + Feedback
    q_keys = [("dq_response_sla", "Response SLA"), ("dq_quality_bar", "Quality bar"), ("dq_audit", "Audit/QA cadence"), ("dq_recovery", "Recovery protocol")]
    q_rows = [(label, d.get(k)) for k, label in q_keys if d.get(k)]
    if q_rows:
        _docx_styled_run(doc.add_paragraph(), "Quality Standards", bold=True, size_pt=12, color=_BRONZE)
        for label, v in q_rows:
            pp = doc.add_paragraph()
            _docx_styled_run(pp, f"{label}: ", bold=True)
            pp.add_run(v)
    f_keys = [("df_nps_cadence", "NPS / Pulse cadence"), ("df_listening", "Listening surface"), ("df_close", "Close-the-loop"), ("df_signal", "Strongest signal to act on")]
    f_rows = [(label, d.get(k)) for k, label in f_keys if d.get(k)]
    if f_rows:
        _docx_styled_run(doc.add_paragraph(), "Feedback Loop", bold=True, size_pt=12, color=_BRONZE)
        for label, v in f_rows:
            pp = doc.add_paragraph()
            _docx_styled_run(pp, f"{label}: ", bold=True)
            pp.add_run(v)
    # AI to-do list
    todos = (_safe_json(d.get("business_plan_todos_json")) or {}).get("todos") or []
    if todos:
        _docx_styled_run(doc.add_paragraph(), "Do This Next — AI-Prioritized 10-Step List", bold=True, size_pt=12, color=_BRONZE)
        for i, t in enumerate(todos[:10]):
            _docx_styled_run(doc.add_paragraph(), f"{i + 1}. {t.get('title', '—')}", bold=True, size_pt=11)
            if t.get("rationale"):
                doc.add_paragraph(t["rationale"])


_DOCX_STEP_RENDERERS = {
    2: _docx_render_step2,
    3: _docx_render_step3,
    4: _docx_render_step4,
    5: _docx_render_step5,
    6: _docx_render_step6,
    7: _docx_render_step7,
}


def _docx_step_extras(doc: Document, step_num: int, d: Dict[str, str]) -> None:
    renderer = _DOCX_STEP_RENDERERS.get(step_num)
    if renderer:
        renderer(doc, d)


# =============================== ENDPOINTS ===============================

async def _load_plan_and_inputs(plan_id: str, user: CurrentUser):
    cli = anon_client_with_token(user.token)
    res = cli.table("plans").select("*").eq("id", plan_id).limit(1).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Plan not found")
    plan = res.data[0]
    inputs = cli.table("plan_inputs").select("*").eq("plan_id", plan_id).execute().data or []
    profile_res = cli.table("profiles").select("*").eq("id", user.id).limit(1).execute()
    profile = profile_res.data[0] if profile_res.data else None
    is_pro = has_pro_access(profile)
    return plan, inputs, is_pro


def _safe_filename(name: str) -> str:
    s = "".join(c if c.isalnum() or c in (" ", "-", "_") else "_" for c in (name or "Plan"))
    return s.strip().replace(" ", "_")[:80] or "Plan"


@router.get("/{plan_id}/export.pdf")
async def export_pdf(plan_id: str, user: CurrentUser = Depends(require_user)):
    plan, inputs, is_pro = await _load_plan_and_inputs(plan_id, user)
    pdf_bytes = _build_pdf(plan, inputs, is_pro)
    filename = f"{_safe_filename(plan.get('title') or 'Plan')}.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/{plan_id}/export.docx")
async def export_docx(plan_id: str, user: CurrentUser = Depends(require_user)):
    plan, inputs, is_pro = await _load_plan_and_inputs(plan_id, user)
    if not is_pro:
        raise HTTPException(status_code=402, detail={"code": "pro_required", "message": "Word export is a Pro feature. Upgrade to download as .docx."})
    docx_bytes = _build_docx(plan, inputs)
    filename = f"{_safe_filename(plan.get('title') or 'Plan')}.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
